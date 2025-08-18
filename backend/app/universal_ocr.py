"""
Universal OCR Service - Extract text from ANY file type
Supports: PDF, images, DOCX, TXT, CSV, Excel
"""

import os
import logging
from typing import Dict, Any, Optional
from pathlib import Path

logger = logging.getLogger(__name__)

class UniversalOCR:
    """Universal text extraction from any file type"""
    
    def __init__(self):
        self.supported_formats = {
            # Text documents
            '.txt', '.md', '.csv',
            # Office documents  
            '.docx', '.doc',
            # PDFs
            '.pdf',
            # Images
            '.png', '.jpg', '.jpeg', '.tiff', '.bmp', '.gif',
            # Spreadsheets
            '.xlsx', '.xls'
        }
    
    def extract_text(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text from any supported file type
        Returns: {'text': content, 'method': extraction_method, 'success': bool}
        """
        if not os.path.exists(file_path):
            return {'text': '', 'method': 'error', 'success': False, 'error': 'File not found'}
        
        file_ext = Path(file_path).suffix.lower()
        
        try:
            if file_ext == '.txt':
                return self._extract_txt(file_path)
            elif file_ext == '.docx':
                return self._extract_docx(file_path)
            elif file_ext == '.pdf':
                return self._extract_pdf(file_path)
            elif file_ext in {'.png', '.jpg', '.jpeg', '.tiff', '.bmp', '.gif'}:
                return self._extract_image(file_path)
            elif file_ext in {'.csv', '.xlsx', '.xls'}:
                return self._extract_data_file(file_path)
            else:
                return {'text': '', 'method': 'unsupported', 'success': False, 'error': f'Unsupported format: {file_ext}'}
                
        except Exception as e:
            logger.error(f"OCR extraction failed for {file_path}: {e}")
            return {'text': '', 'method': 'error', 'success': False, 'error': str(e)}
    
    def _extract_txt(self, file_path: str) -> Dict[str, Any]:
        """Extract text from TXT files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            return {'text': text, 'method': 'direct_read', 'success': True}
        except UnicodeDecodeError:
            # Try different encodings
            for encoding in ['latin-1', 'cp1252', 'utf-16']:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        text = f.read()
                    return {'text': text, 'method': f'direct_read_{encoding}', 'success': True}
                except:
                    continue
            return {'text': '', 'method': 'encoding_error', 'success': False, 'error': 'Could not decode text file'}
    
    def _extract_docx(self, file_path: str) -> Dict[str, Any]:
        """Extract text from DOCX files"""
        try:
            from docx import Document
            doc = Document(file_path)
            
            # Extract paragraphs
            paragraphs = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text)
            
            # Extract tables
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        tables_text.append(' | '.join(row_text))
            
            # Combine all text
            all_text = []
            if paragraphs:
                all_text.extend(paragraphs)
            if tables_text:
                all_text.append('\n--- TABLES ---')
                all_text.extend(tables_text)
            
            final_text = '\n\n'.join(all_text)
            return {'text': final_text, 'method': 'python_docx', 'success': True}
            
        except ImportError:
            return {'text': '', 'method': 'missing_library', 'success': False, 'error': 'python-docx not installed'}
        except Exception as e:
            return {'text': '', 'method': 'docx_error', 'success': False, 'error': str(e)}
    
    def _extract_pdf(self, file_path: str) -> Dict[str, Any]:
        """Extract text from PDF files"""
        text_content = ""
        method = "none"
        
        # Method 1: Try PyPDF2 first (faster)
        try:
            import PyPDF2
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                pages_text = []
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    if page_text.strip():
                        pages_text.append(f"--- Page {page_num + 1} ---\n{page_text}")
                
                if pages_text:
                    text_content = '\n\n'.join(pages_text)
                    method = "PyPDF2"
        except Exception as e:
            logger.warning(f"PyPDF2 failed for {file_path}: {e}")
        
        # Method 2: Try pdfplumber if PyPDF2 fails or gives poor results
        if not text_content or len(text_content.strip()) < 50:
            try:
                import pdfplumber
                with pdfplumber.open(file_path) as pdf:
                    pages_text = []
                    for page_num, page in enumerate(pdf.pages):
                        page_text = page.extract_text()
                        if page_text and page_text.strip():
                            pages_text.append(f"--- Page {page_num + 1} ---\n{page_text}")
                    
                    if pages_text:
                        text_content = '\n\n'.join(pages_text)
                        method = "pdfplumber"
            except ImportError:
                pass
            except Exception as e:
                logger.warning(f"pdfplumber failed for {file_path}: {e}")
        
        if text_content.strip():
            return {'text': text_content, 'method': method, 'success': True}
        else:
            return {'text': '', 'method': 'pdf_no_text', 'success': False, 'error': 'No text found in PDF'}
    
    def _extract_image(self, file_path: str) -> Dict[str, Any]:
        """Extract text from images using OCR"""
        try:
            import pytesseract
            from PIL import Image
            import cv2
            import numpy as np
            
            # Load image
            image = cv2.imread(file_path)
            if image is None:
                # Fallback to PIL
                pil_image = Image.open(file_path)
                image = cv2.cvtColor(np.array(pil_image), cv2.COLOR_RGB2BGR)
            
            # Preprocess image for better OCR
            gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
            
            # Apply denoising
            denoised = cv2.medianBlur(gray, 3)
            
            # Increase contrast
            contrast = cv2.convertScaleAbs(denoised, alpha=1.5, beta=0)
            
            # Apply threshold
            _, binary = cv2.threshold(contrast, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
            
            # Extract text
            custom_config = r'--oem 3 --psm 6'
            text = pytesseract.image_to_string(binary, config=custom_config)
            
            if text.strip():
                return {'text': text.strip(), 'method': 'tesseract_ocr', 'success': True}
            else:
                return {'text': '', 'method': 'ocr_no_text', 'success': False, 'error': 'No text detected in image'}
                
        except ImportError as e:
            return {'text': '', 'method': 'missing_ocr_library', 'success': False, 'error': f'OCR library missing: {e}'}
        except Exception as e:
            return {'text': '', 'method': 'ocr_error', 'success': False, 'error': str(e)}
    
    def _extract_data_file(self, file_path: str) -> Dict[str, Any]:
        """Extract text summary from data files (CSV, Excel)"""
        try:
            import pandas as pd
            
            # Load data file
            if file_path.endswith('.csv'):
                df = pd.read_csv(file_path)
            else:  # Excel
                df = pd.read_excel(file_path)
            
            # Create text summary
            summary_parts = []
            summary_parts.append(f"Dataset Summary: {df.shape[0]} rows, {df.shape[1]} columns")
            summary_parts.append(f"Columns: {', '.join(df.columns.tolist())}")
            
            # Add sample data
            if len(df) > 0:
                summary_parts.append("\nFirst 3 rows:")
                summary_parts.append(df.head(3).to_string())
            
            # Add basic statistics for numeric columns
            numeric_cols = df.select_dtypes(include=['number']).columns
            if len(numeric_cols) > 0:
                summary_parts.append(f"\nNumeric columns statistics:")
                summary_parts.append(df[numeric_cols].describe().to_string())
            
            text_summary = '\n'.join(summary_parts)
            return {'text': text_summary, 'method': 'pandas_summary', 'success': True}
            
        except Exception as e:
            return {'text': '', 'method': 'data_error', 'success': False, 'error': str(e)}
    
    def is_supported(self, file_path: str) -> bool:
        """Check if file type is supported"""
        ext = Path(file_path).suffix.lower()
        return ext in self.supported_formats
    
    def get_content_type(self, file_path: str, extracted_text: str) -> str:
        """Determine content type based on file and text"""
        ext = Path(file_path).suffix.lower()
        
        if ext in {'.csv', '.xlsx', '.xls'}:
            return 'data'
        elif ext == '.docx':
            return 'document'
        elif ext == '.pdf':
            return 'pdf_document'
        elif ext in {'.png', '.jpg', '.jpeg', '.tiff', '.bmp', '.gif'}:
            return 'image_with_text'
        elif ext == '.txt':
            return 'text_file'
        else:
            return 'unknown'

# Global OCR instance
universal_ocr = UniversalOCR()
