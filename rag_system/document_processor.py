"""
DATASOPH AI - Document Processor
Handles document loading, chunking, and preprocessing for RAG system
"""

import os
import re
from typing import List, Dict, Any, Optional
from pathlib import Path
import logging

# Document processing imports
from pypdf import PdfReader
from docx import Document
import pandas as pd

# LangChain imports
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document as LangChainDocument

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentProcessor:
    """
    Document Processor for RAG System
    Handles PDF, DOCX, TXT, and CSV files
    """
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        """
        Initialize Document Processor
        
        Args:
            chunk_size: Size of text chunks
            chunk_overlap: Overlap between chunks
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )
        
        # Supported file types
        self.supported_extensions = {
            '.pdf': self._process_pdf,
            '.docx': self._process_docx,
            '.txt': self._process_txt,
            '.csv': self._process_csv,
            '.json': self._process_json
        }
    
    def process_document(self, file_path: str) -> List[LangChainDocument]:
        """
        Process a document and return chunks
        
        Args:
            file_path: Path to the document
            
        Returns:
            List of LangChain Document chunks
        """
        try:
            file_path = Path(file_path)
            
            if not file_path.exists():
                raise FileNotFoundError(f"File not found: {file_path}")
            
            file_extension = file_path.suffix.lower()
            
            if file_extension not in self.supported_extensions:
                raise ValueError(f"Unsupported file type: {file_extension}")
            
            # Extract text from document
            text = self.supported_extensions[file_extension](file_path)
            
            if not text.strip():
                raise ValueError("Document is empty or could not be processed")
            
            # Split text into chunks
            chunks = self.text_splitter.split_text(text)
            
            # Convert to LangChain Documents
            documents = []
            for i, chunk in enumerate(chunks):
                doc = LangChainDocument(
                    page_content=chunk,
                    metadata={
                        "source": str(file_path),
                        "chunk_id": i,
                        "total_chunks": len(chunks),
                        "file_type": file_extension,
                        "chunk_size": len(chunk)
                    }
                )
                documents.append(doc)
            
            logger.info(f"Processed {file_path.name}: {len(documents)} chunks created")
            return documents
            
        except Exception as e:
            logger.error(f"Error processing document {file_path}: {e}")
            raise
    
    def _process_pdf(self, file_path: Path) -> str:
        """Extract text from PDF file"""
        try:
            reader = PdfReader(file_path)
            text = ""
            
            for page_num, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text.strip():
                    text += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
            
            return text
            
        except Exception as e:
            logger.error(f"Error processing PDF {file_path}: {e}")
            raise
    
    def _process_docx(self, file_path: Path) -> str:
        """Extract text from DOCX file"""
        try:
            doc = Document(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            return text
            
        except Exception as e:
            logger.error(f"Error processing DOCX {file_path}: {e}")
            raise
    
    def _process_txt(self, file_path: Path) -> str:
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            # Try with different encoding
            with open(file_path, 'r', encoding='latin-1') as f:
                return f.read()
        except Exception as e:
            logger.error(f"Error processing TXT {file_path}: {e}")
            raise
    
    def _process_csv(self, file_path: Path) -> str:
        """Extract text from CSV file"""
        try:
            df = pd.read_csv(file_path)
            
            # Convert DataFrame to text representation
            text = f"CSV Document: {file_path.name}\n"
            text += f"Shape: {df.shape}\n"
            text += f"Columns: {list(df.columns)}\n\n"
            
            # Add sample data
            text += "Sample Data:\n"
            text += df.head(10).to_string()
            
            # Add column descriptions
            text += "\n\nColumn Descriptions:\n"
            for col in df.columns:
                text += f"- {col}: {df[col].dtype}\n"
            
            return text
            
        except Exception as e:
            logger.error(f"Error processing CSV {file_path}: {e}")
            raise
    
    def _process_json(self, file_path: Path) -> str:
        """Extract text from JSON file"""
        try:
            import json
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            # Convert JSON to text representation
            text = f"JSON Document: {file_path.name}\n"
            text += json.dumps(data, indent=2, ensure_ascii=False)
            
            return text
            
        except Exception as e:
            logger.error(f"Error processing JSON {file_path}: {e}")
            raise
    
    def process_multiple_documents(self, file_paths: List[str]) -> List[LangChainDocument]:
        """
        Process multiple documents
        
        Args:
            file_paths: List of file paths
            
        Returns:
            List of all document chunks
        """
        all_documents = []
        
        for file_path in file_paths:
            try:
                documents = self.process_document(file_path)
                all_documents.extend(documents)
                logger.info(f"Successfully processed {file_path}")
            except Exception as e:
                logger.error(f"Failed to process {file_path}: {e}")
                continue
        
        logger.info(f"Total documents processed: {len(all_documents)}")
        return all_documents
    
    def get_document_summary(self, documents: List[LangChainDocument]) -> Dict[str, Any]:
        """
        Get summary of processed documents
        
        Args:
            documents: List of processed documents
            
        Returns:
            Summary dictionary
        """
        if not documents:
            return {"error": "No documents provided"}
        
        sources = set()
        total_chunks = len(documents)
        total_chars = sum(len(doc.page_content) for doc in documents)
        
        for doc in documents:
            if "source" in doc.metadata:
                sources.add(doc.metadata["source"])
        
        return {
            "total_documents": len(sources),
            "total_chunks": total_chunks,
            "total_characters": total_chars,
            "average_chunk_size": total_chars / total_chunks if total_chunks > 0 else 0,
            "sources": list(sources),
            "chunk_size": self.chunk_size,
            "chunk_overlap": self.chunk_overlap
        }


def main():
    """Demo function for Document Processor"""
    print("📄 DATASOPH AI - Document Processor Demo")
    print("=" * 50)
    
    # Initialize processor
    processor = DocumentProcessor(chunk_size=1000, chunk_overlap=200)
    
    print(f"Supported file types: {list(processor.supported_extensions.keys())}")
    print(f"Chunk size: {processor.chunk_size}")
    print(f"Chunk overlap: {processor.chunk_overlap}")
    print("-" * 50)
    
    # Example usage
    print("Example usage:")
    print("processor = DocumentProcessor()")
    print("documents = processor.process_document('sample.pdf')")
    print("summary = processor.get_document_summary(documents)")


if __name__ == "__main__":
    main() 